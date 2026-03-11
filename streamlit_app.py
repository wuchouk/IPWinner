import re
import streamlit as st
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.drawing.spreadsheet_drawing import TwoCellAnchor, AnchorMarker
from openpyxl.drawing.image import Image
from openpyxl.utils.units import pixels_to_EMU
from io import BytesIO
from datetime import datetime, timezone, timedelta
import streamlit.components.v1 as st_components
import os, tempfile, json, subprocess
import urllib.request
import urllib.error
import base64
import ssl
import zipfile
import time
import hashlib

APP_VERSION = "v13"


def _get_git_commit_utc():
    """從 git log 取得最新 commit 的 UTC 時間（datetime 物件）"""
    try:
        result = subprocess.run(
            ['git', 'log', '-1', '--format=%ci'],
            capture_output=True, text=True, timeout=5,
            cwd=os.path.dirname(__file__),
        )
        if result.returncode == 0 and result.stdout.strip():
            # git 輸出格式：2026-03-04 19:32:32 +0000
            raw = result.stdout.strip()
            dt = datetime.strptime(raw, '%Y-%m-%d %H:%M:%S %z')
            return dt.astimezone(timezone.utc)
    except Exception:
        pass
    return None


_GIT_COMMIT_UTC = _get_git_commit_utc()

# ============================================================
# 專利公開資訊 API 模組（台灣案全自動下載）
# ============================================================
_TIPO_API_BASE = "https://tiponet.tipo.gov.tw/S092_API/opd1"
_SSL_CTX = ssl.create_default_context()


def tipo_get_token(account, password):
    """取得 Bearer Token（專利公開資訊 API）"""
    creds = base64.b64encode(f"{account}:{password}".encode()).decode()
    req = urllib.request.Request(f"{_TIPO_API_BASE}/getAuth")
    req.add_header("Authorization", f"Basic {creds}")
    with urllib.request.urlopen(req, timeout=30, context=_SSL_CTX) as resp:
        return resp.read().decode("utf-8").strip()


def _strip_tw_prefix(patent_id):
    """去除 TW 前綴和 A/B 後綴，保留 I/M/D 前綴。
    例：TW202421040A → 202421040, TWI740797 → I740797, TW I756159 → I756159"""
    s = patent_id.strip()
    # 去掉 TW 前綴（含可能的空格）
    if s.upper().startswith("TW"):
        s = s[2:].lstrip()
    # 去掉尾部的 A1/A2/B1/B2 等後綴（但不影響 I/M/D 開頭的號碼）
    s = re.sub(r'[AB]\d*$', '', s)
    return s


def tipo_get_case_info(token, case_id):
    """查詢專利案件基本資訊（可用申請案號、公開號、公告號）。
    自動去除 TW 前綴。回傳 dict（新版 API 直接回傳欄位，不包裹 caseInformation）。"""
    clean_id = _strip_tw_prefix(case_id)
    req = urllib.request.Request(f"{_TIPO_API_BASE}/getCaseInfo/{clean_id}")
    req.add_header("Authorization", f"Bearer {token}")
    with urllib.request.urlopen(req, timeout=30, context=_SSL_CTX) as resp:
        return json.loads(resp.read().decode("utf-8"))


def tipo_get_file_list(token, case_no):
    """查詢收發文歷程（含說明書下載連結）"""
    req = urllib.request.Request(f"{_TIPO_API_BASE}/getResultFileList/{case_no}")
    req.add_header("Authorization", f"Bearer {token}")
    with urllib.request.urlopen(req, timeout=30, context=_SSL_CTX) as resp:
        return json.loads(resp.read().decode("utf-8"))


def tipo_download_file(token, file_url):
    """下載檔案（串流），回傳 bytes"""
    req = urllib.request.Request(file_url)
    req.add_header("Authorization", f"Bearer {token}")
    with urllib.request.urlopen(req, timeout=120, context=_SSL_CTX) as resp:
        return resp.read()


def tipo_find_latest_specification(file_list_data):
    """從收發文歷程中找到最新版的專利說明書 PDF。
    搜尋策略：從後往前找 showName 包含「說明書」的檔案。
    回傳 (showName, fileURL, length) 或 None。"""
    results = file_list_data.get("resultFileList", [])
    # 從最新的紀錄往回找
    for record in reversed(results):
        for f in record.get("fileList", []):
            name = f.get("showName", "")
            if "說明書" in name and f.get("fileURL"):
                return {
                    "showName": name,
                    "fileURL": f["fileURL"],
                    "length": f.get("length", 0),
                    "category": record.get("caseReasonName", ""),
                    "date": record.get("documentDate", ""),
                }
    return None


# ============================================================
# GPSS API 模組（外國案查詢 + 連結產生）
# ============================================================
_GPSS_API_BASE = "https://tiponet.tipo.gov.tw/gpss1/gpsskmc/gpss_api"
_GPSS_SEARCH_URL = "https://tiponet.tipo.gov.tw/gpss2/gpsskmc/gpssbkm"

# 資料庫代碼對照（完整版，依據 GPSS API v1.4 官方文件）
# A=公開/申請, B=公告/核准, D=設計
_COUNTRY_DB_MAP = {
    "TW": ["TWA", "TWB", "TWD"],
    "US": ["USA", "USB", "USD"],
    "CN": ["CNA", "CNB", "CND"],
    "JP": ["JPA", "JPB", "JPD"],
    "EP": ["EPA", "EPB", "EUIPO"],
    "KR": ["KPA", "KPB", "KPD"],
    "WO": ["WO"],
    "SE": ["SEAA", "SEAB"],      # 東南亞（含東協各國）
    "OT": ["OTA", "OTB"],        # 其他國家
}

# 反向：從 DB 代碼推回國家
_DB_TO_COUNTRY = {}
for _c, _dbs in _COUNTRY_DB_MAP.items():
    for _d in _dbs:
        _DB_TO_COUNTRY[_d] = _c


def gpss_search(user_code, patent_number, pat_db=None, exp_fld=None):
    """用 GPSS API 查詢專利（用 PN 或 AN）。回傳 JSON dict 或 None。"""
    params = f"userCode={user_code}"
    if pat_db:
        params += f"&patDB={pat_db}"
    params += f"&PN={patent_number}"
    # 預設輸出欄位：公告號、申請號、公告日、名稱、申請人、發明人
    fields = exp_fld or "PN,ID,AN,AD,TI,PA,IN"
    params += f"&expFld={fields}&expFmt=json&expQty=5"
    url = f"{_GPSS_API_BASE}?{params}"
    req = urllib.request.Request(url)
    req.add_header("User-Agent", "Mozilla/5.0")
    try:
        with urllib.request.urlopen(req, timeout=30, context=_SSL_CTX) as resp:
            data = json.loads(resp.read().decode("utf-8"))
            return data
    except Exception:
        return None


def gpss_verify_patent(user_code, country, number):
    """
    用 GPSS API 驗證專利是否存在，並回傳 API 提供的精確公告號碼。
    回傳 dict: {"found": bool, "doc_number": str, "title": str, "db": str}
    精確的 doc_number 可用來建構更準確的 Google Patents URL。
    """
    result = {"found": False, "doc_number": "", "title": "", "db": ""}
    dbs = _COUNTRY_DB_MAP.get(country, [])
    if not dbs:
        return result

    for db in dbs:
        data = gpss_search(user_code, number, pat_db=db)
        if not data:
            continue
        try:
            # 解析 JSON 結構：gpss-API.patent.patentcontent[]
            patent_node = data.get("gpss-API", {}).get("patent", {})
            contents = patent_node.get("patentcontent", [])
            if isinstance(contents, dict):
                contents = [contents]
            if not contents:
                continue
            pc = contents[0]
            pub_ref = pc.get("publication-reference", {})
            doc_num = pub_ref.get("doc-number", "")
            title_node = pc.get("patent-title", "")
            if isinstance(title_node, dict):
                title_node = title_node.get("#text", "")
            if doc_num:
                result["found"] = True
                result["doc_number"] = doc_num
                result["title"] = title_node
                result["db"] = db
                return result
        except Exception:
            continue

    return result


# ============================================================
# 外國案直接連結產生模組
# ============================================================

def _build_foreign_patent_links(country, number):
    """
    根據國家碼和專利號產生各國專利資料庫的直接連結。
    回傳 list of dict: [{"source": "來源名稱", "url": "連結"}]
    """
    links = []
    # 清理號碼：去除國碼前綴，保留純數字（及可能的小數點）
    clean_num = number.strip()

    if country == "US":
        # Google Patents
        links.append({
            "source": "Google Patents",
            "url": f"https://patents.google.com/patent/US{clean_num}",
        })
        # USPTO Full-Text (application)
        # 格式判斷：如果是 7 位以上純數字 → 授權號；如果含 "/" 或以 "20" 開頭 → 申請號
        digits_only = clean_num.replace(",", "").replace(" ", "")
        if "/" in clean_num:
            # 申請號格式 e.g. 16/123,456
            links.append({
                "source": "USPTO",
                "url": f"https://patft.uspto.gov/netacgi/nph-Parser?Sect1=PTO1&Sect2=HITOFF&p=1&u=/netahtml/PTO/srchnum.htm&r=1&f=G&l=50&d=PALL&s1={clean_num}.PN.",
            })
        else:
            links.append({
                "source": "USPTO",
                "url": f"https://patft.uspto.gov/netacgi/nph-Parser?Sect1=PTO1&Sect2=HITOFF&p=1&u=/netahtml/PTO/srchnum.htm&r=1&f=G&l=50&d=PALL&s1={digits_only}.PN.",
            })

    elif country == "CN":
        # Google Patents（最穩定的中國專利 PDF 來源）
        links.append({
            "source": "Google Patents",
            "url": f"https://patents.google.com/patent/CN{clean_num}",
        })

    elif country == "JP":
        # Google Patents
        links.append({
            "source": "Google Patents",
            "url": f"https://patents.google.com/patent/JP{clean_num}",
        })
        # J-PlatPat
        links.append({
            "source": "J-PlatPat",
            "url": f"https://www.j-platpat.inpit.go.jp/c1801/PU/JP-{clean_num}/11/ja",
        })

    elif country == "EP":
        # Espacenet（歐洲專利局，可直接看 PDF）
        links.append({
            "source": "Espacenet",
            "url": f"https://worldwide.espacenet.com/patent/search?q=pn%3DEP{clean_num}",
        })
        # Google Patents
        links.append({
            "source": "Google Patents",
            "url": f"https://patents.google.com/patent/EP{clean_num}",
        })

    elif country == "KR":
        # Google Patents
        links.append({
            "source": "Google Patents",
            "url": f"https://patents.google.com/patent/KR{clean_num}",
        })
        # KIPRIS（韓國智慧財產局）
        links.append({
            "source": "KIPRIS",
            "url": f"https://kpat.kipris.or.kr/kpat/searchLogina.do?next=MainSearch&lng=en",
        })

    elif country == "WO":
        # Google Patents
        links.append({
            "source": "Google Patents",
            "url": f"https://patents.google.com/patent/WO{clean_num}",
        })
        # WIPO PATENTSCOPE
        links.append({
            "source": "PATENTSCOPE",
            "url": f"https://patentscope.wipo.int/search/en/detail.jsf?docId=WO{clean_num}",
        })

    else:
        # 其他國家：fallback 到 Google Patents
        links.append({
            "source": "Google Patents",
            "url": f"https://patents.google.com/patent/{country}{clean_num}",
        })

    # 所有外國案都附上 GPSS 作為備用來源
    links.append({
        "source": "GPSS",
        "url": "https://tiponet.tipo.gov.tw/gpss2/gpsskmc/gpssbkm",
    })

    return links


# ============================================================
# 專利號碼解析模組
# ============================================================
# 各國專利號碼格式（帶國碼前綴）
_PATENT_PATTERNS = [
    # TW: TW105131793, TWI642307, TW201637458A
    (r'\b(TW)\s*([A-Z]?\d{6,}[A-Z]?\d*)\b', 'TW'),
    # US: US20150001234A1, US9876543B2, US2015/0001234
    (r'\b(US)\s*(\d{4,}/?[\d]+[A-Z]?\d*)\b', 'US'),
    # CN: CN201510879928A, CN1234567B, ZL202210112039.7（ZL 為中國授權專利前綴）
    (r'\b(CN)\s*(\d{6,}[A-Z]?\d*)\b', 'CN'),
    (r'\b(ZL)\s*(\d{6,}[\d.]*)\b', 'CN'),  # ZL 開頭 → 視為 CN
    # JP: JP2015123456A
    (r'\b(JP)\s*(\d{4,}[A-Z]?\d*)\b', 'JP'),
    # EP: EP1234567A1
    (r'\b(EP)\s*(\d{4,}[A-Z]?\d*)\b', 'EP'),
    # KR: KR20150001234A
    (r'\b(KR)\s*(\d{4,}[A-Z]?\d*)\b', 'KR'),
    # WO: WO2015001234A1
    (r'\b(WO)\s*(\d{4,}[A-Z]?\d*)\b', 'WO'),
]

# 純數字號碼（不帶國碼）
_BARE_NUMBER_PATTERN = re.compile(r'\b(\d{6,}[A-Z]?\d*)\b')


def parse_patent_numbers(text):
    """從文字中解析專利號碼。
    回傳 list of dict: [{"country": "TW"|""|..., "number": "...", "raw": "原始文字"}]
    """
    results = []
    seen = set()

    # 先找帶國碼的
    found_positions = set()  # 避免同一位置被重複匹配
    for pattern, country in _PATENT_PATTERNS:
        for m in re.finditer(pattern, text, re.IGNORECASE):
            raw = m.group(0).strip()
            prefix = m.group(1).upper()
            number = m.group(2).replace("/", "").replace(" ", "")
            # ZL 開頭的保留完整號碼（含小數點），顯示時加回 ZL 前綴
            if prefix == "ZL":
                display_number = f"ZL{number}"
            else:
                display_number = number
            key = f"{country}_{number}"
            if key not in seen:
                seen.add(key)
                results.append({"country": country, "number": display_number, "raw": raw})
                found_positions.update(range(m.start(), m.end()))

    # 收集所有已識別的號碼（不含國碼前綴），用於 bare number dedup
    _known_numbers = set()
    for r in results:
        # 取出純數字部分用於比對
        _n = r["number"].replace("ZL", "").replace(".", "")
        _known_numbers.add(_n)

    # 再找純數字（沒有國碼前綴的）
    for m in _BARE_NUMBER_PATTERN.finditer(text):
        # 跳過已經被帶國碼模式匹配到的位置
        if any(p in found_positions for p in range(m.start(), m.end())):
            continue
        number = m.group(1)
        # 跳過已被其他國碼模式匹配到的相同號碼
        if number in _known_numbers:
            continue
        key = f"__{number}"
        if key not in seen:
            seen.add(key)
            results.append({"country": "", "number": number, "raw": number})

    return results


def parse_file_for_patent_numbers(uploaded_file):
    """從上傳的檔案（txt/docx/xlsx）中提取專利號碼"""
    filename = uploaded_file.name.lower()
    text = ""

    if filename.endswith('.txt'):
        text = uploaded_file.read().decode('utf-8', errors='replace')
    elif filename.endswith(('.doc', '.docx')):
        try:
            from docx import Document
            doc = Document(BytesIO(uploaded_file.read()))
            text = "\n".join(p.text for p in doc.paragraphs)
            # 也讀取表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
        except ImportError:
            # 若 python-docx 未安裝，嘗試純文字讀取
            text = uploaded_file.read().decode('utf-8', errors='replace')
    elif filename.endswith(('.xlsx', '.xls')):
        wb = openpyxl.load_workbook(BytesIO(uploaded_file.read()), data_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        text += str(cell) + "\n"
    else:
        text = uploaded_file.read().decode('utf-8', errors='replace')

    return parse_patent_numbers(text)


# ============================================================
# API 金鑰加密儲存模組
# ============================================================
_CRED_FILE = os.path.join(os.path.dirname(__file__), '.api_credentials.json')
_ENC_KEY = hashlib.sha256(b"IPWinner2026SecretKey").digest()


def _encrypt_str(text):
    """簡易加密（XOR + base64）"""
    encrypted = bytes([b ^ _ENC_KEY[i % len(_ENC_KEY)] for i, b in enumerate(text.encode('utf-8'))])
    return base64.b64encode(encrypted).decode('ascii')


def _decrypt_str(enc_text):
    """簡易解密"""
    data = base64.b64decode(enc_text)
    decrypted = bytes([b ^ _ENC_KEY[i % len(_ENC_KEY)] for i, b in enumerate(data)])
    return decrypted.decode('utf-8')


def _get_user_id():
    """取得使用者識別碼（Streamlit Cloud 有登入時用 email hash，否則用 'default'）"""
    try:
        user = st.experimental_user
        if user and hasattr(user, 'email') and user.email:
            return hashlib.md5(user.email.encode()).hexdigest()
    except Exception:
        pass
    return "default"


def _load_api_credentials():
    """讀取當前使用者的 API 帳密"""
    try:
        with open(_CRED_FILE, 'r') as f:
            data = json.load(f)
        uid = _get_user_id()
        if uid in data:
            return {
                'account': _decrypt_str(data[uid]['account']),
                'password': _decrypt_str(data[uid]['password']),
            }
    except Exception:
        pass
    return None


def _save_api_credentials(account, password):
    """儲存當前使用者的 API 帳密（加密）"""
    try:
        with open(_CRED_FILE, 'r') as f:
            data = json.load(f)
    except Exception:
        data = {}
    uid = _get_user_id()
    data[uid] = {
        'account': _encrypt_str(account),
        'password': _encrypt_str(password),
    }
    with open(_CRED_FILE, 'w') as f:
        json.dump(data, f)


def _load_secrets_fallback():
    """從 st.secrets 讀取預設 API 帳密（Streamlit Cloud 後台設定，不會因重新部署消失）"""
    try:
        return {
            'account': st.secrets["tipo_api"]["account"],
            'password': st.secrets["tipo_api"]["password"],
        }
    except Exception:
        return None


def _get_api_credentials():
    """取得 API 帳密（優先使用者儲存 → fallback st.secrets）。
    回傳 (account, password, source) 或 (None, None, None)。
    source: 'user' | 'default' | None"""
    saved = _load_api_credentials()
    if saved:
        return saved['account'], saved['password'], 'user'
    fallback = _load_secrets_fallback()
    if fallback:
        return fallback['account'], fallback['password'], 'default'
    return None, None, None


# ============================================================
# 頁面設定
# ============================================================
st.set_page_config(
    page_title="IP Winner 工具箱",
    page_icon="📋",
    layout="wide",
)

# 自訂 CSS：拖放上傳區域放大 + 視覺提示 + 檔案列表一次顯示更多
st.markdown("""
<style>
/* ---- 放大拖放上傳區域 ---- */
[data-testid="stFileUploaderDropzone"] {
    min-height: 220px;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    border: 2px dashed #b0b8c8;
    border-radius: 12px;
    background: #f8fafd;
    transition: border-color 0.2s, background 0.2s;
}
[data-testid="stFileUploaderDropzone"]:hover {
    border-color: #4F8BF9;
    background: #eef3fc;
}
/* 讓拖放區域內的文字和按鈕居中 */
[data-testid="stFileUploaderDropzone"] > div {
    display: flex;
    flex-direction: column;
    align-items: center;
    gap: 4px;
}
/* 隱藏 file_uploader 分頁按鈕 */
[data-testid="stFileUploader"] nav[role="navigation"],
[data-testid="stFileUploader"] [data-testid="stPagination"],
[data-testid="stFileUploader"] .stPagination {
    display: none !important;
}
/* 檔案列表容器高度 */
[data-testid="stFileUploader"] [data-testid="stFileUploaderFileList"] {
    max-height: 700px !important;
    overflow-y: auto !important;
}
</style>
""", unsafe_allow_html=True)

# ============================================================
# 偵測使用者瀏覽器時區（用於下載檔名時間戳）
# ============================================================
@st.cache_resource
def _create_tz_component():
    """建立瀏覽器時區偵測元件（使用 Streamlit declare_component）"""
    comp_dir = os.path.join(tempfile.gettempdir(), "_st_tz_detect")
    os.makedirs(comp_dir, exist_ok=True)
    with open(os.path.join(comp_dir, "index.html"), "w") as f:
        f.write("""<!DOCTYPE html>
<html><body><script>
function send(type, data) {
    window.parent.postMessage(
        Object.assign({isStreamlitMessage: true, type: type}, data), "*"
    );
}
send("streamlit:componentReady", {apiVersion: 1});
window.addEventListener("message", function(event) {
    if (event.data.type === "streamlit:render") {
        send("streamlit:setComponentValue",
             {value: new Date().getTimezoneOffset()});
    }
});
</script></body></html>""")
    return st_components.declare_component("_tz_detect", path=comp_dir)

_tz_component = _create_tz_component()
_client_tz_offset = _tz_component(default=0, key="_tz_offset", height=0)


def _get_client_now():
    """取得使用者本地時間（根據瀏覽器時區）"""
    if isinstance(_client_tz_offset, (int, float)):
        # getTimezoneOffset() 回傳「UTC - 本地」的分鐘數，所以要取反
        client_tz = timezone(timedelta(minutes=-int(_client_tz_offset)))
        return datetime.now(client_tz)
    return datetime.now(timezone.utc)


# ============================================================
# 合併紀錄持久化（JSON 檔案）
# ============================================================
_HISTORY_FILE = os.path.join(os.path.dirname(__file__), '.merge_history.json')


def _load_merge_history():
    """從 JSON 檔讀取今日合併紀錄；若非今日或檔案不存在則回傳空 list"""
    try:
        with open(_HISTORY_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if data.get('date') == _get_client_now().strftime('%Y-%m-%d'):
            return data.get('records', [])
    except (FileNotFoundError, json.JSONDecodeError, ValueError):
        pass
    return []


def _save_merge_history(records):
    """將合併紀錄寫入 JSON 檔（含今日日期）"""
    data = {
        'date': _get_client_now().strftime('%Y-%m-%d'),
        'records': records,
    }
    try:
        with open(_HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except OSError:
        pass  # Streamlit Cloud 等環境可能寫入失敗，靜默處理


# ============================================================
# 合併檔標題檔案解析
# ============================================================
_HEADER_FILENAME_RE = re.compile(r'.+-([^-]+)-合併檔標題')


def parse_header_filename(filename):
    """從標題檔檔名提取慧盈案號。
    格式：日期-XX-案號-合併檔標題.xlsx，案號為「合併檔標題」前一段。
    例如 'yyyymmdd-IP-KOIS23004WWW1-合併檔標題.xlsx' → 'KOIS23004WWW1'
    回傳 (案號, 錯誤訊息)，其中一個為 None。"""
    m = _HEADER_FILENAME_RE.search(filename)
    if not m:
        return None, "標題檔檔名需包含「合併檔標題」，格式：日期-XX-案號-合併檔標題.xlsx"
    return m.group(1), None


def read_header_file(file_bytes):
    """讀取合併檔標題檔案，提取 Row 1 LOGO、Row 2 文字資訊，以及監控商標圖片。
    圖片分類邏輯：col >= 4 為 LOGO（右側）、col < 4 為商標圖片（左側）。
    回傳 dict，含 logo_*、trademark_*、row2_* 等欄位。"""
    wb = openpyxl.load_workbook(BytesIO(file_bytes))
    ws = wb.active
    result = {
        'logo_image_data': None,
        'logo_width': None,
        'logo_height': None,
        'trademark_images': [],   # 監控商標圖片（可能有 0~多張）
        'row2_values': {},
        'row2_merges': [],
        'row1_height': ws.row_dimensions[1].height or 58,
        'row2_height': ws.row_dimensions[2].height or 38,
    }

    # 讀取 Row 1~2 範圍內的所有圖片，按位置分類
    for img in ws._images:
        anchor = img.anchor
        if not hasattr(anchor, '_from'):
            continue
        # 只處理 Row 0~1 的圖片（對應 Excel Row 1~2）
        if anchor._from.row > 1:
            continue

        img_data = BytesIO(img._data())
        from_info = {
            'col': anchor._from.col,
            'colOff': anchor._from.colOff,
            'row': anchor._from.row,
            'rowOff': anchor._from.rowOff,
        }

        if anchor._from.col >= 4:
            # 右側 → LOGO
            result['logo_image_data'] = img_data
            result['logo_width'] = img.width
            result['logo_height'] = img.height
        else:
            # 左側 → 監控商標圖片
            to_info = None
            if hasattr(anchor, 'to') and anchor.to:
                to_info = {
                    'col': anchor.to.col,
                    'colOff': anchor.to.colOff,
                    'row': anchor.to.row,
                    'rowOff': anchor.to.rowOff,
                }
            result['trademark_images'].append({
                'data': img_data,
                'width': img.width,
                'height': img.height,
                'from': from_info,
                'to': to_info,
            })

    # 讀取 Row 2 合併儲存格範圍
    from openpyxl.utils import get_column_letter
    for mc in ws.merged_cells.ranges:
        if mc.min_row == 2 and mc.max_row == 2:
            result['row2_merges'].append(str(mc))

    # 讀取 Row 2 的 cell 資料（只讀有值的 cell）
    for c in range(1, 13):
        cell = ws.cell(row=2, column=c)
        if cell.value is not None:
            cl = get_column_letter(c)
            result['row2_values'][cl] = {
                'value': cell.value,
                'font_name': cell.font.name,
                'font_size': cell.font.size,
                'font_bold': cell.font.bold,
                'align_h': cell.alignment.horizontal,
                'align_v': cell.alignment.vertical,
                'align_wrap': cell.alignment.wrap_text,
            }

    wb.close()
    return result


# ============================================================
# 欄位對應設定
# ============================================================
DB1_MAPPING = {
    'A': 'B', 'B': 'C', 'C': 'D', 'D': 'E', 'E': 'F',
    'F': 'G', 'G': 'H', 'H': 'I', 'I': 'J', 'J': 'K', 'K': 'L',
}
DB1_IMAGE_SRC_COL = 1

# DB2 有兩種格式：
# 格式 A（有地區欄）：A:序号, B:商标文字, C:商标图样, D:地區, E:类别, F:申请号, G:申请人, H:申请日期, I:初审公告日期, J:初审公告期号, K:异议截止日期, L:群组, M:商品/服务
# 格式 B（無地區欄）：A:序号, B:商标文字, C:商标图样, D:类别, E:申请号, F:申请人, G:申请日期, H:初审公告日期, I:初审公告期号, J:异议截止日期, K:群组, L:商品/服务
DB2_MAPPING_WITH_REGION = {
    'B': 'B',   # 商标文字 → 他人商標
    'C': 'C',   # 商标图样 → 商標圖樣
    'D': 'D',   # 地區 → 地區
    'E': 'E',   # 类别 → 商標類別
    'F': 'F',   # 申请号 → 申請號
    'G': 'G',   # 申请人 → 申請人
    'H': 'H',   # 申请日期 → 申請日期
    'I': 'I',   # 初审公告日期 → 公告日期
    'K': 'J',   # 异议截止日期 → 異議期限
    'M': 'K',   # 商品/服务 → 商品/服務名稱（原文）
}
DB2_MAPPING_NO_REGION = {
    'B': 'B',   # 商标文字 → 他人商標
    'C': 'C',   # 商标图样 → 商標圖樣
    'D': 'E',   # 类别 → 商標類別
    'E': 'F',   # 申请号 → 申請號
    'F': 'G',   # 申请人 → 申請人
    'G': 'H',   # 申请日期 → 申請日期
    'H': 'I',   # 初审公告日期 → 公告日期
    'J': 'J',   # 异议截止日期 → 異議期限
    'L': 'K',   # 商品/服务 → 商品/服務名稱（原文）
}
DB2_IMAGE_SRC_COL = 2


def _detect_db2_has_region(ws):
    """偵測 DB2 工作表是否含有「地區」欄位（在 Row 1 的 D 欄）"""
    headers = _get_row1_headers(ws)
    return '地區' in headers

DB3_MAPPING = {
    'B': 'B', 'C': 'C',
    'D': 'D', 'E': 'E', 'F': 'F', 'G': 'G', 'H': 'H', 'I': 'I', 'J': 'J',
    'K': 'K', 'L': 'L',
}
DB3_IMAGE_SRC_COL = 2

# 合併檔（已合併過的檔案，可以再次與新資料合併）
MERGED_FILE_MAPPING = {
    'B': 'B', 'C': 'C', 'D': 'D', 'E': 'E', 'F': 'F',
    'G': 'G', 'H': 'H', 'I': 'I', 'J': 'J', 'K': 'K', 'L': 'L',
}
MERGED_FILE_IMAGE_SRC_COL = 2
MERGED_FILE_DATA_START = 3  # 合併檔的資料從第 3 列開始

# DB3 需要清理的欄位（去除前綴、轉換日期格式）
MONTH_MAP = {
    'JAN': '01', 'FEB': '02', 'MAR': '03', 'APR': '04',
    'MAY': '05', 'JUN': '06', 'JUL': '07', 'AUG': '08',
    'SEP': '09', 'OCT': '10', 'NOV': '11', 'DEC': '12',
}


def _convert_ddmonyyyy(text):
    """DD-MON-YYYY → YYYY-MM-DD，例如 '02-JAN-2026' → '2026-01-02'。
    無法辨識則回傳原文。"""
    m = re.match(r'^(\d{1,2})-([A-Z]{3})-(\d{4})$', text.strip())
    if m:
        day, mon, year = m.groups()
        month_num = MONTH_MAP.get(mon)
        if month_num:
            return f'{year}-{month_num}-{day.zfill(2)}'
    return text


def clean_db3_date(value):
    """將 DB3 日期格式轉為 YYYY-MM-DD，例如 'Reg: 22-AUG-2025' → '2025-08-22'。
    不管前綴是什麼（App: / Reg: / Opp 等）都去除，只保留日期並轉換格式。"""
    if not value or not isinstance(value, str):
        return value
    value = value.strip()
    # 去除任何英文前綴（含冒號和空格），只保留日期部分
    value = re.sub(r'^[A-Za-z]+[:\s]*\s*', '', value)
    return _convert_ddmonyyyy(value)


def clean_db3_opposition(value):
    """清理異議期限（J 欄）：去 Opp 前綴，並將所有 DD-MMM-YYYY 轉為 YYYY-MM-DD。
    保留其餘文字不動。
    例如 'Opp CN : 30-APR-2026\\nCA : 2 months...' → 'CN : 2026-04-30\\nCA : 2 months...'"""
    if not value or not isinstance(value, str):
        return value
    text = value.strip()
    # 去除開頭的 "Opp " 前綴
    text = re.sub(r'^Opp\s+', '', text)
    # 替換所有出現的 DD-MON-YYYY
    def _replace_date(m):
        return _convert_ddmonyyyy(m.group(0))
    text = re.sub(r'\d{1,2}-[A-Z]{3}-\d{4}', _replace_date, text)
    return text


def clean_class_column(value):
    """清理商標類別（E 欄）：去空格、頓號分隔、去前導 0。
    例如 '09 11 43' → '9、11、43'，'9, 11, 43' → '9、11、43'"""
    if not value:
        return value
    text = str(value).strip()
    if not text:
        return text
    # 以逗號、頓號、空白拆分
    parts = re.split(r'[,、\s]+', text)
    # 去前導 0 並過濾空字串
    cleaned = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # 去前導 0（純數字才處理）
        if p.isdigit():
            p = str(int(p))
        cleaned.append(p)
    return '、'.join(cleaned) if cleaned else text


def clean_db3_app_number(value):
    """去除申請號的前綴，例如 'App 26000442' → '26000442'"""
    if not value or not isinstance(value, str):
        return value
    return re.sub(r'^(App|Reg)\s+', '', value.strip())


# DB3 地區名稱對照表
DB3_REGION_MAP = {
    'EU trade marks': 'EUIPO',
    'International Register': 'WIPO',
    'United States of America': 'US (USPTO)',
}


def clean_db3_region(value):
    """轉換 DB3 地區名稱，例如 'EU trade marks' → 'EUIPO'"""
    if not value or not isinstance(value, str):
        return value
    value = value.strip()
    # 先檢查完全匹配
    if value in DB3_REGION_MAP:
        return DB3_REGION_MAP[value]
    # 處理帶括號的情況，例如 "EU trade marks (unpublished applications)" → "EUIPO(unpublished applications)"
    for original, replacement in DB3_REGION_MAP.items():
        if value.startswith(original):
            suffix = value[len(original):]  # 例如 " (unpublished applications)"
            return replacement + suffix.lstrip()  # "EUIPO(unpublished applications)"
    return value

MERGED_HEADERS = [
    '#', '他人商標', '商標圖樣', '地區', '商標類別',
    '申請號', '申請人', '申請日期', '公告日期', '異議期限',
    '商品/服務名稱（原文）', '商品/服務名稱（英文翻譯）',
]
MERGED_IMAGE_COL = 2
MERGED_HEADER_ROW = 3
MERGED_DATA_START = 4
SOURCE_DATA_START = 2


# ============================================================
# 核心函式
# ============================================================
def col_letter_to_index(letter):
    result = 0
    for ch in letter.upper():
        result = result * 26 + (ord(ch) - ord('A') + 1)
    return result - 1


def find_merged_header_row(file_bytes):
    """找出合併檔的標頭所在列（掃描 Row 1~5）"""
    try:
        wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
        ws = wb.active
        for check_row in range(1, min(6, ws.max_row + 1)):
            headers = set()
            for cell in ws[check_row]:
                if cell.value is not None:
                    headers.add(str(cell.value).strip())
            if '他人商標' in headers and '商標圖樣' in headers:
                wb.close()
                return check_row
        wb.close()
    except Exception:
        pass
    return None


def _get_row1_headers(ws):
    """取得工作表 Row 1 的所有非空值"""
    headers = set()
    for cell in ws[1]:
        if cell.value is not None:
            headers.add(str(cell.value).strip())
    return headers


def detect_db_type(file_bytes):
    """
    自動辨識檔案來自哪個資料庫。
    回傳 'db1', 'db2', 'db3', 'merged', 或 None（無法辨識）。

    注意：某些 xlsx 在 read_only 模式會讀不到完整資料，
    因此改用一般模式；DB1 (Markify) 的資料可能不在第一個 sheet，
    需要掃描所有 sheet。
    """
    try:
        wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)

        # 先檢查所有 sheet（DB1 Markify 資料在非第一個 sheet）
        for ws in wb.worksheets:
            headers_row1 = _get_row1_headers(ws)

            # DB1：英文標頭，含 "Trademark"
            if 'Trademark' in headers_row1:
                wb.close()
                return 'db1'

            # DB2：簡體中文，含 "商标文字"
            if '商标文字' in headers_row1:
                wb.close()
                return 'db2'

            # DB3：繁體中文，含 "他人商標"（Row 1 有 # 和 他人商標）
            if '他人商標' in headers_row1:
                wb.close()
                return 'db3'

        # 都不符合 → 檢查 active sheet 的 Row 2~5（合併檔）
        ws = wb.active
        for check_row in range(2, min(6, ws.max_row + 1)):
            headers_check = set()
            for cell in ws[check_row]:
                if cell.value is not None:
                    headers_check.add(str(cell.value).strip())
            if '他人商標' in headers_check and '商標圖樣' in headers_check:
                wb.close()
                return 'merged'

        wb.close()
        return None
    except Exception:
        return None


def _find_data_sheets(wb, db_type):
    """找出含有資料的 sheet，掃描所有 tab 避免資料不在 active sheet 的問題"""
    # 各 DB 的 Row 1 關鍵字
    db_signature = {
        'db1': 'Trademark',
        'db2': '商标文字',
        'db3': '他人商標',
    }
    keyword = db_signature.get(db_type)
    if keyword:
        sheets = []
        for ws in wb.worksheets:
            headers = _get_row1_headers(ws)
            if keyword in headers:
                sheets.append(ws)
        return sheets if sheets else [wb.active]
    # merged 或其他：使用 active sheet
    return [wb.active]


def read_source_data(file_bytes, mapping, db_type=''):
    """讀取來源 Excel 並按 mapping 轉換欄位"""
    wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
    rows = []

    sheets = _find_data_sheets(wb, db_type)

    for ws in sheets:
        # DB2：動態偵測該 sheet 是否含有「地區」欄位，選用對應的 mapping
        if db_type == 'db2':
            if _detect_db2_has_region(ws):
                active_mapping = DB2_MAPPING_WITH_REGION
                db2_default_region = None
            else:
                active_mapping = DB2_MAPPING_NO_REGION
                db2_default_region = 'China'  # 摩知輪本身是中國資料庫
        else:
            active_mapping = mapping
            db2_default_region = None

        # 合併檔需要動態找到標頭列，資料從標頭列 +1 開始
        if db_type == 'merged':
            header_row = find_merged_header_row(file_bytes) or 2
            data_start = header_row + 1
        else:
            data_start = SOURCE_DATA_START
        for row_idx, row in enumerate(
            ws.iter_rows(min_row=data_start, values_only=False),
            start=data_start,
        ):
            if row_idx > ws.max_row:
                break
            merged_row = {}
            for src_col_letter, dest_col_letter in active_mapping.items():
                src_idx = col_letter_to_index(src_col_letter)
                if src_idx < len(row):
                    merged_row[dest_col_letter] = row[src_idx].value
            # DB2：無地區欄時預設填入 China（摩知輪 = 中國資料庫）
            if db_type == 'db2' and db2_default_region:
                if not merged_row.get('D') or str(merged_row['D']).strip() == '':
                    merged_row['D'] = db2_default_region
            # DB3 資料清理
            if db_type == 'db3':
                if 'F' in merged_row:
                    merged_row['F'] = clean_db3_app_number(merged_row['F'])
                if 'D' in merged_row:
                    merged_row['D'] = clean_db3_region(merged_row['D'])
                for date_col in ['H', 'I']:
                    if date_col in merged_row:
                        merged_row[date_col] = clean_db3_date(merged_row[date_col])
                if 'J' in merged_row:
                    merged_row['J'] = clean_db3_opposition(merged_row['J'])
                # 空的公告日期 → 1900-01-00、空的異議期限 → 0
                if not merged_row.get('I') or str(merged_row['I']).strip() == '':
                    merged_row['I'] = '1900-01-00'
                if not merged_row.get('J') or str(merged_row['J']).strip() == '':
                    merged_row['J'] = '0'
            if any(v is not None for v in merged_row.values()):
                rows.append(merged_row)
    wb.close()
    return rows


def read_source_images(file_bytes, src_image_col, db_type=''):
    """讀取來源檔案中的圖片"""
    wb = openpyxl.load_workbook(BytesIO(file_bytes))
    images = {}
    sheets = _find_data_sheets(wb, db_type)
    row_offset_accum = 0  # 多 sheet 時累積行數偏移

    for sheet_idx, ws in enumerate(sheets):
        for img in ws._images:
            anchor = img.anchor
            if hasattr(anchor, '_from') and anchor._from.col == src_image_col:
                src_row = anchor._from.row + row_offset_accum
                img_data = BytesIO(img._data())
                images[src_row] = {
                    'data': img_data,
                    'width': img.width,
                    'height': img.height,
                    'from_colOff': anchor._from.colOff,
                    'from_rowOff': anchor._from.rowOff,
                    'to_col': anchor.to.col,
                    'to_row': anchor.to.row,
                    'to_colOff': anchor.to.colOff,
                    'to_rowOff': anchor.to.rowOff,
                    # 保留原始跨距（DB2 圖片會跨到下一列/欄）
                    'row_span': anchor.to.row - anchor._from.row,
                    'col_span': anchor.to.col - anchor._from.col,
                }
        # 累積偏移：當前 sheet 的資料行數（不含標頭）
        if sheet_idx < len(sheets) - 1:
            row_offset_accum += ws.max_row - SOURCE_DATA_START + 1
    wb.close()
    return images


def create_merged_file(all_rows, all_images, header_data=None, progress_bar=None):
    """建立合併後的 Excel 檔。header_data 來自 read_header_file()。"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '商標監控結果清單'

    # 使用 scheme='minor' 讓字型參照 theme 設定（英文 Times New Roman / 中文 新細明體）
    header_font = Font(scheme='minor', bold=True, size=11)
    header_fill = PatternFill('solid', fgColor='D9E1F2')
    header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin'),
    )
    data_font = Font(scheme='minor', size=10)
    align_center_top = Alignment(horizontal='center', vertical='top', wrap_text=True)
    align_left_top = Alignment(horizontal='left', vertical='top', wrap_text=True)

    # ── Row 1：LOGO（從標題檔載入，水平＋垂直置中） ──
    ws.merge_cells('A1:L1')
    row1_h = (header_data or {}).get('row1_height', 58)
    ws.row_dimensions[1].height = row1_h
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')

    if header_data and header_data.get('logo_image_data'):
        header_data['logo_image_data'].seek(0)
        logo_img = Image(header_data['logo_image_data'])
        logo_img.width = header_data.get('logo_width', 863)
        logo_img.height = header_data.get('logo_height', 133)
        # 縮放到合理高度（行高 - 3pt 留白）
        target_h = max(int(row1_h - 3), 40)
        if logo_img.height > target_h:
            ratio = target_h / logo_img.height
            logo_img.height = target_h
            logo_img.width = int(logo_img.width * ratio)

        # 計算各欄寬度（像素）以求出水平置中偏移
        _cw_chars = [4, 23, 23, 15, 15, 16, 25, 15, 20, 20, 60, 60]  # A~L
        _cw_px = [int(w * 7 + 5) for w in _cw_chars]
        _total_px = sum(_cw_px)
        _left_px = (_total_px - logo_img.width) / 2

        _cum = 0
        _logo_col, _logo_col_off_px = 0, 0
        for _i, _cpx in enumerate(_cw_px):
            if _cum + _cpx > _left_px:
                _logo_col = _i
                _logo_col_off_px = int(_left_px - _cum)
                break
            _cum += _cpx

        _row_h_emu = int(row1_h * 12700)
        _img_h_emu = pixels_to_EMU(logo_img.height)
        _row_off_emu = max(0, (_row_h_emu - _img_h_emu) // 2)

        _right_px = _left_px + logo_img.width
        _cum2 = 0
        _logo_to_col, _logo_to_col_off_px = len(_cw_px) - 1, _cw_px[-1]
        for _i, _cpx in enumerate(_cw_px):
            if _cum2 + _cpx >= _right_px:
                _logo_to_col = _i
                _logo_to_col_off_px = int(_right_px - _cum2)
                break
            _cum2 += _cpx

        _from_marker = AnchorMarker(
            col=_logo_col,
            colOff=pixels_to_EMU(_logo_col_off_px),
            row=0,
            rowOff=_row_off_emu,
        )
        _to_marker = AnchorMarker(
            col=_logo_to_col,
            colOff=pixels_to_EMU(_logo_to_col_off_px),
            row=0,
            rowOff=_row_off_emu + _img_h_emu,
        )
        logo_img.anchor = TwoCellAnchor(_from=_from_marker, to=_to_marker)
        ws.add_image(logo_img)

    # ── Row 2：監控商標 / 監控地區 / 監控類別（從標題檔載入） ──
    if header_data and header_data.get('row2_merges'):
        for mr in header_data['row2_merges']:
            ws.merge_cells(mr)
    else:
        ws.merge_cells('A2:C2')
        ws.merge_cells('D2:G2')
        ws.merge_cells('H2:L2')
    row2_h = (header_data or {}).get('row2_height', 38)
    ws.row_dimensions[2].height = row2_h

    if header_data and header_data.get('row2_values'):
        for cl, info in header_data['row2_values'].items():
            ci = col_letter_to_index(cl) + 1
            cell = ws.cell(row=2, column=ci, value=info['value'])
            cell.font = Font(
                name=info.get('font_name', '新細明體'),
                size=info.get('font_size', 14),
                bold=info.get('font_bold', True),
            )
            cell.alignment = Alignment(
                horizontal=info.get('align_h', 'left'),
                vertical=info.get('align_v', 'center'),
                wrap_text=info.get('align_wrap', True),
            )
    else:
        # 備用：沒有標題檔時使用預設空白標題
        info_font = Font(scheme='minor', bold=True, size=14)
        info_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
        for col_letter, label in [('A', '監控商標：'), ('D', '監控地區：'), ('H', '監控類別：')]:
            ci = col_letter_to_index(col_letter) + 1
            cell = ws.cell(row=2, column=ci, value=label)
            cell.font = info_font
            cell.alignment = info_align

    # ── 監控商標圖片（底部對齊 Row 2 底線上方 5pt，向上延伸到 Row 1） ──
    if header_data and header_data.get('trademark_images'):
        _PT = 12700  # 1pt = 12700 EMU
        _row1_h_emu = int(row1_h * _PT)
        _row2_h_emu = int(row2_h * _PT)
        _margin_emu = int(5 * _PT)  # 底線往上 5pt
        for tm_img_info in header_data['trademark_images']:
            tm_img_info['data'].seek(0)
            tm_img = Image(tm_img_info['data'])
            tm_img.width = tm_img_info['width']
            tm_img.height = tm_img_info['height']
            f = tm_img_info['from']
            # 計算原始圖片顯示高度（EMU）
            if tm_img_info.get('to'):
                t = tm_img_info['to']
                # absolute_y = 各 row 高度加總 + rowOff
                _orig_from_y = f['rowOff']
                for _r in range(f['row']):
                    _orig_from_y += _row1_h_emu if _r == 0 else _row2_h_emu
                _orig_to_y = t['rowOff']
                for _r in range(t['row']):
                    _orig_to_y += _row1_h_emu if _r == 0 else _row2_h_emu
                _img_h_emu = max(_orig_to_y - _orig_from_y, int(20 * _PT))
            else:
                _img_h_emu = pixels_to_EMU(tm_img.height)
            # to = Row 2 底線往上 5pt
            _to_abs_y = _row1_h_emu + _row2_h_emu - _margin_emu
            # from = to 往上 圖片高度
            _from_abs_y = max(0, _to_abs_y - _img_h_emu)
            # 轉換回 row / rowOff
            if _from_abs_y < _row1_h_emu:
                _from_row, _from_off = 0, _from_abs_y
            else:
                _from_row, _from_off = 1, _from_abs_y - _row1_h_emu
            if _to_abs_y < _row1_h_emu:
                _to_row, _to_off = 0, _to_abs_y
            else:
                _to_row, _to_off = 1, _to_abs_y - _row1_h_emu
            _tm_from = AnchorMarker(
                col=f['col'], colOff=f['colOff'],
                row=_from_row, rowOff=int(_from_off),
            )
            _to_col = t['col'] if tm_img_info.get('to') else f['col']
            _to_colOff = t['colOff'] if tm_img_info.get('to') else (
                f['colOff'] + pixels_to_EMU(tm_img.width))
            _tm_to = AnchorMarker(
                col=_to_col, colOff=_to_colOff,
                row=_to_row, rowOff=int(_to_off),
            )
            tm_img.anchor = TwoCellAnchor(_from=_tm_from, to=_tm_to)
            ws.add_image(tm_img)

    # ── Row 3：欄位標題 ──
    for col_idx, header in enumerate(MERGED_HEADERS, start=1):
        cell = ws.cell(row=MERGED_HEADER_ROW, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    total = len(all_rows)
    col_letters = ['B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L']
    for i, row_data in enumerate(all_rows):
        row_num = MERGED_DATA_START + i
        num_cell = ws.cell(row=row_num, column=1, value=f'=ROW()-{MERGED_HEADER_ROW}')
        num_cell.font = data_font
        num_cell.alignment = align_center_top
        num_cell.border = thin_border
        for cl in col_letters:
            ci = col_letter_to_index(cl) + 1
            v = row_data.get(cl)
            # 申請號（F 欄）：強制存為文字
            if cl == 'F' and v is not None:
                v = str(v).strip()
            # 商標類別（E 欄）：去空格、頓號分隔、去前導 0
            if cl == 'E' and v is not None:
                v = clean_class_column(v)
            cell = ws.cell(row=row_num, column=ci, value=v if v is not None else '')
            cell.font = data_font
            cell.alignment = align_left_top
            cell.border = thin_border
        if progress_bar and i % 50 == 0:
            progress_bar.progress(
                0.3 + 0.3 * (i / max(total, 1)),
                text=f'寫入資料 {i}/{total}...',
            )

    col_widths = {
        'A': 4, 'B': 23, 'C': 23, 'D': 15, 'E': 15,
        'F': 16, 'G': 25, 'H': 15, 'I': 20, 'J': 20, 'K': 60, 'L': 60,
    }
    for cl, w in col_widths.items():
        ws.column_dimensions[cl].width = w

    for r in range(MERGED_DATA_START, MERGED_DATA_START + total):
        ws.row_dimensions[r].height = 100

    img_total = len(all_images)
    for idx, (img_info, row_offset) in enumerate(all_images):
        try:
            img_info['data'].seek(0)
            img = Image(img_info['data'])
            if img_info['width'] and img_info['height']:
                img.width = img_info['width']
                img.height = img_info['height']
            else:
                img.width = 80
                img.height = 80
            max_size = 100
            if img.width > max_size or img.height > max_size:
                ratio = min(max_size / img.width, max_size / img.height)
                img.width = int(img.width * ratio)
                img.height = int(img.height * ratio)
            new_row = img_info['orig_row'] + row_offset
            new_col = MERGED_IMAGE_COL
            # 圖片頂部至少留 100000 EMU (~7.6pt) 的間距，避免壓到上方框線
            MIN_ROW_OFF = 100000
            orig_row_off = img_info.get('from_rowOff', 0)
            row_off = max(orig_row_off, MIN_ROW_OFF)
            _from = AnchorMarker(
                col=new_col,
                colOff=img_info.get('from_colOff', 0),
                row=new_row,
                rowOff=row_off,
            )
            _to = AnchorMarker(
                col=new_col + img_info.get('col_span', 0),
                colOff=img_info.get('to_colOff', pixels_to_EMU(img.width)),
                row=new_row + img_info.get('row_span', 0),
                rowOff=img_info.get('to_rowOff', pixels_to_EMU(img.height)),
            )
            img.anchor = TwoCellAnchor(_from=_from, to=_to)
            ws.add_image(img)
        except Exception:
            pass
        if progress_bar and idx % 20 == 0:
            progress_bar.progress(
                0.6 + 0.35 * (idx / max(img_total, 1)),
                text=f'寫入圖片 {idx}/{img_total}...',
            )

    ws.freeze_panes = 'A4'
    last_row = MERGED_HEADER_ROW + len(all_rows)
    ws.auto_filter.ref = f'A{MERGED_HEADER_ROW}:L{last_row}'

    # ── 列印設定（需求 7~10） ──
    # 需求 7：頁尾 — 頁碼/總頁數（置中）
    ws.oddFooter.center.text = "&P / &N"
    # 需求 8：列印標題列（每頁重複 Row 1~3）
    ws.print_title_rows = '1:3'
    # 需求 9：頁面邊界 — 上下左右 2cm (≈0.787in)，頁首頁尾 1cm (≈0.394in)，水平置中
    ws.page_margins.top = 0.787
    ws.page_margins.bottom = 0.787
    ws.page_margins.left = 0.787
    ws.page_margins.right = 0.787
    ws.page_margins.header = 0.394
    ws.page_margins.footer = 0.394
    ws.print_options.horizontalCentered = True
    # 需求 10：列印方向橫向，自動縮放至 1 頁寬（取代固定 46%）
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.paperSize = 9  # A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0  # 高度不限頁數
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    # Patch theme XML：將 minor font 的 latin 字型改為 Times New Roman
    # 這樣 scheme='minor' 的 cell 會用 Times New Roman 顯示英文，新細明體顯示中文
    from zipfile import ZipFile
    patched = BytesIO()
    with ZipFile(output, 'r') as zin:
        with ZipFile(patched, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'xl/theme/theme1.xml':
                    text = data.decode('utf-8')
                    text = re.sub(
                        r'(<a:minorFont>\s*<a:latin typeface=")[^"]*(")',
                        r'\1Times New Roman\2',
                        text,
                    )
                    data = text.encode('utf-8')
                zout.writestr(item, data)
    patched.seek(0)
    return patched, len(all_rows)


# ============================================================
# Streamlit 介面
# ============================================================
DB_LABELS = {
    'merged': '合併檔',
    'db1': '資料庫 1',
    'db2': '資料庫 2',
    'db3': '資料庫 3',
}

# ============================================================
# Sidebar 導覽
# ============================================================
_PAGES = ["📋 合併檔案", "📥 下載公開說明書", "⚙️ 設定"]

if 'current_page' not in st.session_state:
    st.session_state.current_page = _PAGES[0]

with st.sidebar:
    st.title("📋 IP Winner")
    st.divider()

    # sidebar 按鈕導覽樣式
    st.markdown("""
    <style>
    section[data-testid="stSidebar"] button[kind="secondary"] {
        background: none !important;
        border: none !important;
        box-shadow: none !important;
        text-align: left !important;
        justify-content: flex-start !important;
        padding: 0.5rem 0.75rem !important;
        border-radius: 0.5rem !important;
        width: 100% !important;
        font-size: 1rem !important;
    }
    section[data-testid="stSidebar"] button[kind="secondary"]:hover {
        background-color: rgba(151, 166, 195, 0.15) !important;
    }
    /* 覆蓋按鈕內層 div 的 justify-content: center */
    section[data-testid="stSidebar"] button[kind="secondary"] > div {
        justify-content: flex-start !important;
    }
    </style>
    """, unsafe_allow_html=True)

    for _p in _PAGES:
        _is_active = (st.session_state.current_page == _p)
        _label = f"**{_p}**" if _is_active else _p
        if st.button(_label, key=f"nav_{_p}", use_container_width=True):
            st.session_state.current_page = _p
            st.rerun()

_page = st.session_state.current_page

# ============================================================
# 輔助函式：重置查詢 / 儲存歷史
# ============================================================
def _reset_patent_query():
    """清除當前查詢狀態並重置輸入框"""
    for k in ['patent_download_done', 'patent_results', 'patent_files', 'patent_parsed']:
        st.session_state.pop(k, None)
    st.session_state.patent_input_key = st.session_state.get('patent_input_key', 0) + 1

def _save_to_history():
    """將當前查詢結果存入歷史紀錄"""
    results = st.session_state.get('patent_results')
    files = st.session_state.get('patent_files')
    if not results:
        return
    if 'patent_history' not in st.session_state:
        st.session_state.patent_history = []
    entry = {
        'timestamp': _get_client_now().strftime('%Y-%m-%d %H:%M'),
        'results': results,
        'files': files or {},
        'ok_count': sum(1 for r in results if r['status'] == 'ok'),
        'link_count': sum(1 for r in results if r['status'] == 'link'),
        'total': len(results),
    }
    st.session_state.patent_history.insert(0, entry)

# ============================================================
# 頁面 1: 合併檔案（原有功能）
# ============================================================
if _page == "📋 合併檔案":
    st.title("📋 合併檔案")
    st.markdown("上傳各資料庫的原始 Excel 檔（最多 15 個），系統會自動辨識來源並合併。也可同時放入舊的合併檔，系統會一起整合。")

    st.divider()

    # 用遞增 key 來重置 file_uploader
    if 'uploader_key' not in st.session_state:
        st.session_state.uploader_key = 0

    # ── 上傳區 1：合併檔標題（必填） ──
    st.subheader("① 合併檔標題")
    header_file = st.file_uploader(
        "上傳合併檔標題檔案（必填）",
        type=["xlsx"],
        accept_multiple_files=False,
        help="檔名格式：yyyymmdd-IP-慧盈案號-合併檔標題.xlsx",
        key=f"header_uploader_{st.session_state.uploader_key}",
    )

    # 驗證標題檔
    _header_ok = False
    _header_data = None
    _case_number = None
    if header_file:
        _case_number, _header_err = parse_header_filename(header_file.name)
        if _header_err:
            st.error(f"⚠️ {_header_err}\n\n上傳的檔名：`{header_file.name}`")
        else:
            _header_data = read_header_file(header_file.getvalue())
            _header_ok = True
            st.success(f"✅ 慧盈案號：**{_case_number}**")
    else:
        st.info("請先上傳合併檔標題檔案，才能進行合併。")

    st.divider()

    # ── 上傳區 2：資料檔案 ──
    st.subheader("② 資料檔案")
    uploaded_files = st.file_uploader(
        "將檔案拖放至此處，或點擊 Browse files 選擇檔案",
        type=["xlsx"],
        accept_multiple_files=True,
        help="支援同時上傳多個 .xlsx 檔案，系統會自動辨識來自哪個資料庫",
        key=f"file_uploader_{st.session_state.uploader_key}",
    )

    # 注入 JS 強制顯示所有已上傳檔案（移除 Streamlit 內建的分頁隱藏）
    # st_components.html 的 iframe 有 allow-same-origin，可以存取 parent document
    if uploaded_files and len(uploaded_files) > 3:
        st_components.html("""<script>
        (function(){
            try {
                var doc = window.parent.document;
                function showAll() {
                    var items = doc.querySelectorAll(
                        '[data-testid="stFileUploaderFile"]'
                    );
                    items.forEach(function(el){ el.style.display = 'flex'; });
                    var pag = doc.querySelectorAll(
                        '[data-testid="stFileUploader"] nav[role="navigation"], ' +
                        '[data-testid="stFileUploader"] [data-testid="stPagination"]'
                    );
                    pag.forEach(function(el){ el.style.display = 'none'; });
                }
                showAll();
                setTimeout(showAll, 500);
                setTimeout(showAll, 1500);
            } catch(e) { /* sandbox 擋住就算了，CSS 會處理分頁隱藏 */ }
        })();
        </script>""", height=0)

    # 檢查上傳數量
    if uploaded_files and len(uploaded_files) > 15:
        st.error("⚠️ 最多只能上傳 15 個檔案，請減少檔案數量。")
        st.stop()

    # 辨識並分類檔案
    if uploaded_files:
        classified = {'merged': [], 'db1': [], 'db2': [], 'db3': []}
        unknown_files = []

        for f in uploaded_files:
            file_bytes = f.getvalue()
            db_type = detect_db_type(file_bytes)
            if db_type:
                classified[db_type].append((f.name, file_bytes))
            else:
                unknown_files.append(f.name)

        # 顯示辨識結果
        st.subheader("📂 檔案辨識結果")

        cols = st.columns(4)
        for i, (db_key, label) in enumerate(DB_LABELS.items()):
            with cols[i]:
                files_list = classified[db_key]
                count = len(files_list)
                if count > 0:
                    st.success(f"**{label}**　{count} 個檔案")
                    for fname, _ in files_list:
                        st.caption(f"　📄 {fname}")
                else:
                    if db_key == 'merged':
                        st.info(f"**{label}**　無")
                    else:
                        st.warning(f"**{label}**　未偵測到")

        if unknown_files:
            st.error(
                f"⚠️ 以下 {len(unknown_files)} 個檔案無法辨識來源，將被忽略：\n\n"
                + "\n".join(f"- {name}" for name in unknown_files)
            )

        # 確認至少有檔案可以合併
        total_files = sum(len(v) for v in classified.values())
        if total_files == 0:
            st.error("沒有可辨識的檔案，請確認上傳的是正確的原始檔。")
            st.stop()

        st.divider()

        # 初始化 session_state
        if 'merge_done' not in st.session_state:
            st.session_state.merge_done = False

        # 合併按鈕（僅在尚未合併時顯示，且必須有標題檔）
        if not st.session_state.merge_done:
            _can_merge = _header_ok and total_files > 0
            if st.button("🚀 開始合併", type="primary", use_container_width=True, disabled=not _can_merge):
                progress_bar = st.progress(0, text="開始處理...")
                logs = []

                try:
                    all_rows = []
                    all_images = []
                    row_counts = {}
                    img_counts = {}

                    # 定義處理順序和對應的設定
                    db_configs = {
                        'merged': {'mapping': MERGED_FILE_MAPPING, 'img_col': MERGED_FILE_IMAGE_SRC_COL},
                        'db1': {'mapping': DB1_MAPPING, 'img_col': DB1_IMAGE_SRC_COL},
                        'db2': {'mapping': DB2_MAPPING_NO_REGION, 'img_col': DB2_IMAGE_SRC_COL},
                        'db3': {'mapping': DB3_MAPPING, 'img_col': DB3_IMAGE_SRC_COL},
                    }

                    step = 0
                    total_steps = total_files * 2  # 每個檔案讀資料 + 讀圖片
                    logs.append(f"合併檔標題 / {header_file.name}")

                    # 按 合併檔 → db1 → db2 → db3 順序處理
                    for db_key in ['merged', 'db1', 'db2', 'db3']:
                        files_list = classified[db_key]
                        if not files_list:
                            continue

                        config = db_configs[db_key]
                        label = DB_LABELS[db_key]
                        db_row_count = 0
                        db_img_count = 0

                        for fname, file_bytes in files_list:
                            # 讀取資料
                            step += 1
                            progress_bar.progress(
                                0.05 + 0.20 * (step / total_steps),
                                text=f'讀取 {label}：{fname}...',
                            )
                            rows = read_source_data(file_bytes, config['mapping'], db_type=db_key)

                            # 讀取圖片
                            step += 1
                            progress_bar.progress(
                                0.05 + 0.20 * (step / total_steps),
                                text=f'讀取 {label} 圖片：{fname}...',
                            )
                            images = read_source_images(file_bytes, config['img_col'], db_type=db_key)
                            logs.append(f"{label} / {fname}：{len(rows)} 筆資料 / {len(images)} 張圖片")

                            # 計算圖片位移
                            if db_key == 'merged':
                                src_header_row = find_merged_header_row(file_bytes) or 2
                                src_data_start = src_header_row + 1
                            else:
                                src_data_start = SOURCE_DATA_START
                            row_offset = (MERGED_DATA_START - src_data_start) + len(all_rows)
                            for src_row, img_info in images.items():
                                img_info['orig_row'] = src_row
                                all_images.append((img_info, row_offset))

                            all_rows.extend(rows)
                            db_row_count += len(rows)
                            db_img_count += len(images)

                        row_counts[db_key] = db_row_count
                        img_counts[db_key] = db_img_count

                    # 建立合併檔
                    progress_bar.progress(0.30, text=f'建立合併檔（{len(all_rows)} 筆，{len(all_images)} 張圖片）...')
                    output_bytes, count = create_merged_file(all_rows, all_images, header_data=_header_data, progress_bar=progress_bar)
                    progress_bar.progress(1.0, text="合併完成！")

                    # 執行記錄彙總
                    logs.append("───────────────────")
                    for db_key, label in DB_LABELS.items():
                        rc = row_counts.get(db_key, 0)
                        ic = img_counts.get(db_key, 0)
                        fc = len(classified[db_key])
                        if fc > 0:
                            logs.append(f"{label}：{fc} 個檔案 → {rc} 筆 / {ic} 張圖片")
                    logs.append(f"合計：{count} 筆 / {len(all_images)} 張圖片")

                    # 儲存結果到 session_state，下載後頁面 rerun 時仍可顯示
                    st.session_state.merge_done = True
                    st.session_state.merge_output = output_bytes.getvalue()
                    st.session_state.merge_count = count
                    st.session_state.merge_img_count = len(all_images)
                    st.session_state.merge_logs = logs
                    client_now = _get_client_now()
                    # 需求 5：yyyymmdd-TC-案號-商標監控結果清單(完整).xlsx
                    if _case_number:
                        st.session_state.merge_filename = f"{client_now.strftime('%Y%m%d')}-TC-{_case_number}-商標監控結果清單(完整).xlsx"
                    else:
                        st.session_state.merge_filename = f"{client_now.strftime('%Y%m%d_%H%M')}_合併檔.xlsx"
                    st.session_state.merge_active_dbs = [
                        (db_key, DB_LABELS[db_key], row_counts.get(db_key, 0), img_counts.get(db_key, 0), len(classified[db_key]))
                        for db_key in DB_LABELS if len(classified[db_key]) > 0
                    ]
                    st.rerun()

                except Exception as e:
                    progress_bar.empty()
                    st.error(f"❌ 合併失敗：{str(e)}")
                    import traceback
                    with st.expander("錯誤詳情"):
                        st.code(traceback.format_exc())

        # 下載後自動重置的 callback（同時寫入合併紀錄到 JSON 檔）
        def _reset_after_download():
            # 從 JSON 檔讀取既有紀錄
            history = _load_merge_history()
            # 記錄本次合併
            record = {
                'time': _get_client_now().strftime('%H:%M'),
                'filename': st.session_state.get('merge_filename', ''),
                'logs': st.session_state.get('merge_logs', []),
                'count': st.session_state.get('merge_count', 0),
                'img_count': st.session_state.get('merge_img_count', 0),
            }
            history.append(record)
            _save_merge_history(history)
            # 清除本次合併的暫存
            for key in ['merge_done', 'merge_output', 'merge_count', 'merge_img_count',
                         'merge_logs', 'merge_filename', 'merge_active_dbs']:
                st.session_state.pop(key, None)
            st.session_state.uploader_key += 1

        # 合併完成後：持久顯示結果（即使下載觸發 rerun 也不會消失）
        if st.session_state.merge_done:
            st.balloons()
            st.success("合併完成！")

            # 顯示各來源統計
            active_dbs = st.session_state.merge_active_dbs
            result_cols = st.columns(len(active_dbs)) if active_dbs else st.columns(1)
            for i, (db_key, label, rc, ic, fc) in enumerate(active_dbs):
                with result_cols[i]:
                    st.metric(f"{label}（{fc} 檔）", f"{rc} 筆", f"{ic} 張圖片")

            st.markdown(f"### 合計：{st.session_state.merge_count} 筆資料 / {st.session_state.merge_img_count} 張圖片")

            # 下載按鈕（點擊後自動重置頁面）
            st.download_button(
                label="⬇️ 下載合併檔",
                data=st.session_state.merge_output,
                file_name=st.session_state.merge_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary",
                use_container_width=True,
                on_click=_reset_after_download,
            )

            # 執行記錄
            with st.expander("📝 執行記錄"):
                for log in st.session_state.merge_logs:
                    st.text(log)

    # ============================================================
    # 合併紀錄（從 JSON 檔讀取，當日有效，隔日自動清空）
    # ============================================================
    _persisted_history = _load_merge_history()
    if _persisted_history:
        st.divider()
        with st.expander(f"📋 今日合併紀錄（{len(_persisted_history)} 次）", expanded=False):
            for i, rec in enumerate(reversed(_persisted_history), 1):
                st.markdown(f"**{i}. {rec['time']}　→　{rec['filename']}**　"
                            f"（{rec['count']} 筆 / {rec['img_count']} 張圖片）")
                for log_line in rec['logs']:
                    if log_line.startswith('──'):
                        break
                    st.caption(f"　　{log_line}")
                if i < len(_persisted_history):
                    st.markdown("---")

# ============================================================
# 頁面 2: 下載公開說明書
# ============================================================
elif _page == "📥 下載公開說明書":
    st.title("📥 下載公開說明書")
    st.markdown("輸入專利號碼（每行一個，或上傳檔案），點擊查詢後系統會自動查詢並下載公開說明書。")

    # 讀取 API 帳密（優先使用者自己的 → fallback 預設）
    _api_account, _api_password, _cred_source = _get_api_credentials()
    if _cred_source == 'user':
        pass  # 使用者已設定，不顯示訊息
    elif _cred_source == 'default':
        st.info("ℹ️ 目前使用預設 API 帳號。如需使用自己的帳號，請至「⚙️ 設定」頁面設定。")
    else:
        st.warning("⚠️ 尚未設定 API 帳號，請先至「⚙️ 設定」頁面儲存 API 帳密。")
        _api_account, _api_password = "", ""

    st.divider()

    # -- 專利號碼輸入 --
    st.subheader("① 輸入專利號碼")

    # 用遞增 key 實現重新查詢時清空輸入框
    if 'patent_input_key' not in st.session_state:
        st.session_state.patent_input_key = 0
    _ik = st.session_state.patent_input_key

    _input_method = st.radio(
        "輸入方式",
        ["直接輸入", "上傳檔案"],
        horizontal=True,
        key=f"patent_input_method_{_ik}",
    )

    if _input_method == "直接輸入":
        _patent_text = st.text_area(
            "請輸入專利號碼（每行一個）",
            height=200,
            placeholder="例如：\nTW105131793\nUS20150001234A1\nCN201510879928A\n104142817",
            key=f"patent_text_input_{_ik}",
        )
    else:
        _patent_file = st.file_uploader(
            "上傳包含專利號碼的檔案",
            type=["txt", "doc", "docx", "xlsx"],
            accept_multiple_files=False,
            key=f"patent_file_uploader_{_ik}",
        )

    st.caption("若號碼未指定國碼，將預設為台灣案處理。若為外國案請加上國碼前綴（TW/US/CN/JP/EP/KR/WO）。中國授權專利可用 ZL 前綴。")

    # -- 查詢按鈕（第一步：解析） --
    _can_query = (_input_method == "直接輸入" and st.session_state.get(f"patent_text_input_{_ik}", "").strip()) or \
                 (_input_method == "上傳檔案" and st.session_state.get(f"patent_file_uploader_{_ik}") is not None)

    if st.button("🔍 查詢", type="primary", use_container_width=True, key="btn_patent_query", disabled=not _can_query):
        if _input_method == "直接輸入":
            _parsed = parse_patent_numbers(st.session_state.get(f"patent_text_input_{_ik}", ""))
        else:
            _parsed = parse_file_for_patent_numbers(st.session_state.get(f"patent_file_uploader_{_ik}"))
        st.session_state.patent_parsed = _parsed
        # 清除之前的下載結果
        for _k in ['patent_download_done', 'patent_results', 'patent_files']:
            st.session_state.pop(_k, None)
        st.rerun()

    # -- 解析結果顯示（點了查詢按鈕之後才顯示） --
    _patent_numbers_raw = st.session_state.get("patent_parsed", [])
    if _patent_numbers_raw:
        st.divider()
        st.subheader("② 解析結果")

        # 檢查是否有無國碼的號碼（需要使用者確認）
        _has_bare = any(p["country"] == "" for p in _patent_numbers_raw)
        _tw_numbers = [p for p in _patent_numbers_raw if p["country"] == "TW"]
        _foreign_numbers = [p for p in _patent_numbers_raw if p["country"] not in ("TW", "")]
        _bare_numbers = [p for p in _patent_numbers_raw if p["country"] == ""]

        # 統計
        _stats_cols = st.columns(4)
        with _stats_cols[0]:
            st.metric("總計", f"{len(_patent_numbers_raw)} 筆")
        with _stats_cols[1]:
            st.metric("🇹🇼 台灣", f"{len(_tw_numbers)} 筆")
        with _stats_cols[2]:
            st.metric("🌍 外國", f"{len(_foreign_numbers)} 筆")
        with _stats_cols[3]:
            st.metric("⚠️ 未指定國碼", f"{len(_bare_numbers)} 筆")

        # 顯示清單
        if _tw_numbers:
            with st.expander(f"🇹🇼 台灣專利（{len(_tw_numbers)} 筆）— 可自動下載", expanded=True):
                for p in _tw_numbers:
                    st.text(f"  {p['raw']}")

        if _foreign_numbers:
            with st.expander(f"🌍 外國專利（{len(_foreign_numbers)} 筆）— 產生各國資料庫連結", expanded=True):
                for p in _foreign_numbers:
                    st.text(f"  {p['country']} {p['number']}")

        if _bare_numbers:
            st.warning("⚠️ 以下號碼未指定國碼，將預設為台灣案處理。若為外國案請加上國碼前綴（TW/US/CN/JP/EP/KR/WO）。", icon="⚠️")
            with st.expander(f"⚠️ 未指定國碼（{len(_bare_numbers)} 筆）", expanded=True):
                for p in _bare_numbers:
                    st.text(f"  {p['number']}")

        st.divider()

        # -- 執行下載 --
        st.subheader("③ 下載公開說明書")

        # 初始化 session state
        if 'patent_download_done' not in st.session_state:
            st.session_state.patent_download_done = False

        if not st.session_state.patent_download_done:
            if st.button("🚀 開始查詢與下載", type="primary", use_container_width=True, key="btn_patent_download"):
                _progress = st.progress(0, text="準備中...")
                _results = []  # list of dict: {number, country, status, filename, data, error, gpss_link}
                _downloaded_files = {}  # filename -> bytes

                # 合併清單：台灣案 + 無國碼案（預設台灣）
                _all_tw = _tw_numbers + _bare_numbers
                _total_tw = len(_all_tw)
                _total_foreign = len(_foreign_numbers)

                # -- 步驟 1：取得 TIPO API Token --
                _token = None
                if _all_tw:
                    _progress.progress(0.02, text="取得 API Token...")
                    try:
                        _token = tipo_get_token(_api_account, _api_password)
                    except Exception as e:
                        st.error(f"❌ 無法取得 API Token：{e}")
                        _token = None

                # -- 步驟 2：處理台灣案 --
                _total_steps = _total_tw + _total_foreign
                for idx, pat in enumerate(_all_tw):
                    _num = pat["number"]
                    _pct = 0.05 + 0.85 * (idx / max(_total_steps, 1))
                    _progress.progress(_pct, text=f"下載台灣案 {_num}... ({idx+1}/{_total_tw})")

                    result = {
                        "number": _num,
                        "country": "TW",
                        "raw": pat["raw"],
                        "status": "error",
                        "filename": "",
                        "data": None,
                        "error": "",
                        "gpss_link": "",
                    }

                    if not _token:
                        result["error"] = "無 API Token"
                        result["status"] = "error"
                        _results.append(result)
                        continue

                    try:
                        # 先查案件資訊（API 會自動去除 TW 前綴）
                        case_info = tipo_get_case_info(_token, _num)
                        case_no = None
                        # 新版 API 直接回傳欄位（不包裹 caseInformation）
                        if case_info and case_info.get("code") == "00":
                            case_no = case_info.get("caseNo", "").replace("-", "")
                        # 舊版 API 相容
                        elif case_info and "caseInformation" in case_info:
                            info = case_info["caseInformation"]
                            case_no = info.get("applicationNo", "").replace("-", "")
                        if not case_no:
                            case_no = _strip_tw_prefix(_num)  # fallback

                        time.sleep(1.5)  # rate limiting 保護

                        # 查檔案清單
                        file_list = tipo_get_file_list(_token, case_no)
                        spec = tipo_find_latest_specification(file_list)

                        if spec:
                            time.sleep(1.5)  # rate limiting 保護
                            # 下載說明書
                            _progress.progress(_pct + 0.02, text=f"下載 {_num} 說明書...")
                            pdf_bytes = tipo_download_file(_token, spec["fileURL"])
                            fname = f"TW_{_num}_{spec['showName']}"
                            if not fname.lower().endswith('.pdf'):
                                fname += '.pdf'
                            # 清理檔名中的非法字元
                            fname = re.sub(r'[\\/:*?"<>|]', '_', fname)
                            _downloaded_files[fname] = pdf_bytes
                            result["status"] = "ok"
                            result["filename"] = fname
                            result["data"] = pdf_bytes
                        else:
                            result["status"] = "not_found"
                            result["error"] = "未找到說明書檔案"

                    except urllib.error.HTTPError as e:
                        if e.code == 403:
                            # Rate limiting — 等待後重試一次
                            time.sleep(5)
                            try:
                                case_info = tipo_get_case_info(_token, _num)
                                case_no = case_info.get("caseNo", "").replace("-", "") if case_info and case_info.get("code") == "00" else _strip_tw_prefix(_num)
                                time.sleep(2)
                                file_list = tipo_get_file_list(_token, case_no)
                                spec = tipo_find_latest_specification(file_list)
                                if spec:
                                    time.sleep(2)
                                    pdf_bytes = tipo_download_file(_token, spec["fileURL"])
                                    fname = f"TW_{_num}_{spec['showName']}"
                                    if not fname.lower().endswith('.pdf'):
                                        fname += '.pdf'
                                    fname = re.sub(r'[\\/:*?"<>|]', '_', fname)
                                    _downloaded_files[fname] = pdf_bytes
                                    result["status"] = "ok"
                                    result["filename"] = fname
                                    result["data"] = pdf_bytes
                                else:
                                    result["status"] = "not_found"
                                    result["error"] = "未找到說明書檔案"
                            except Exception as e2:
                                result["status"] = "error"
                                result["error"] = f"重試失敗：{e2}"
                        else:
                            result["status"] = "error"
                            result["error"] = str(e)
                    except Exception as e:
                        result["status"] = "error"
                        result["error"] = str(e)

                    _results.append(result)
                    time.sleep(1.5)  # rate limiting 保護

                # -- 步驟 3：處理外國案（GPSS API 驗證 + 各國專利資料庫直接連結） --
                _gpss_user_code = "963bED2F36842DCD"
                for idx, pat in enumerate(_foreign_numbers):
                    _num = pat["number"]
                    _country = pat["country"]
                    _pct = 0.05 + 0.85 * ((_total_tw + idx) / max(_total_steps, 1))
                    _progress.progress(_pct, text=f"驗證 {_country} {_num}... ({idx+1}/{_total_foreign})")

                    # 用 GPSS API 驗證專利並取得精確的 doc-number
                    _verified = gpss_verify_patent(_gpss_user_code, _country, _num)
                    # 如果 API 回傳了精確號碼，用它來產生更準確的連結
                    _link_num = _num
                    _api_title = ""
                    if _verified["found"] and _verified["doc_number"]:
                        _link_num = _verified["doc_number"]
                        _api_title = _verified["title"]

                    # 產生各國專利資料庫直接連結
                    patent_links = _build_foreign_patent_links(_country, _link_num)
                    _results.append({
                        "number": _num,
                        "country": _country,
                        "raw": pat["raw"],
                        "status": "link",
                        "filename": "",
                        "data": None,
                        "error": "",
                        "patent_links": patent_links,
                        "gpss_link": patent_links[0]["url"] if patent_links else "",
                        "verified": _verified["found"],
                        "api_doc_number": _verified.get("doc_number", ""),
                        "api_title": _api_title,
                    })

                _progress.progress(1.0, text="完成！")

                # 儲存結果到 session_state
                st.session_state.patent_download_done = True
                st.session_state.patent_results = _results
                st.session_state.patent_files = _downloaded_files
                st.rerun()

        # -- 顯示結果 --
        if st.session_state.patent_download_done:
            _results = st.session_state.get("patent_results", [])
            _downloaded_files = st.session_state.get("patent_files", {})

            _ok = [r for r in _results if r["status"] == "ok"]
            _not_found = [r for r in _results if r["status"] == "not_found"]
            _errors = [r for r in _results if r["status"] == "error"]
            _links = [r for r in _results if r["status"] == "link"]

            st.success(f"查詢完成！")
            _res_cols = st.columns(4)
            with _res_cols[0]:
                st.metric("✅ 已下載", f"{len(_ok)} 筆")
            with _res_cols[1]:
                st.metric("🔗 外國案連結", f"{len(_links)} 筆")
            with _res_cols[2]:
                st.metric("⚠️ 未找到", f"{len(_not_found)} 筆")
            with _res_cols[3]:
                st.metric("❌ 錯誤", f"{len(_errors)} 筆")

            # 詳細結果清單
            if _ok:
                with st.expander(f"✅ 已下載（{len(_ok)} 筆）", expanded=True):
                    for r in _ok:
                        st.markdown(f"- **{r['raw']}** → `{r['filename']}`")

            if _links:
                with st.expander(f"🔗 外國案連結（{len(_links)} 筆）— 點擊前往各國專利資料庫", expanded=True):
                    for r in _links:
                        _pl = r.get("patent_links", [])
                        _badge = "✅" if r.get("verified") else "🔍"
                        _title_str = ""
                        if r.get("api_title"):
                            _title_str = f" — {r['api_title'][:60]}"
                        if _pl:
                            _link_parts = " ｜ ".join(
                                f"[{lnk['source']}]({lnk['url']})" for lnk in _pl
                            )
                            st.markdown(f"- {_badge} **{r['country']} {r['number']}**{_title_str} → {_link_parts}")
                        else:
                            st.markdown(f"- {_badge} **{r['country']} {r['number']}**{_title_str} → `{r['number']}`")

            if _not_found:
                with st.expander(f"⚠️ 未找到說明書（{len(_not_found)} 筆）", expanded=False):
                    for r in _not_found:
                        st.markdown(f"- **{r['raw']}**：{r['error']}")

            if _errors:
                with st.expander(f"❌ 查詢失敗（{len(_errors)} 筆）", expanded=False):
                    for r in _errors:
                        st.markdown(f"- **{r['raw']}**：{r['error']}")

            # 下載 ZIP + 重新查詢（並排）
            if _downloaded_files:
                zip_buf = BytesIO()
                with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for fname, fdata in _downloaded_files.items():
                        zf.writestr(fname, fdata)
                zip_buf.seek(0)

                _btn_col1, _btn_col2 = st.columns([2, 1])
                with _btn_col1:
                    st.download_button(
                        label=f"⬇️ 下載全部說明書（{len(_downloaded_files)} 個 PDF，ZIP 壓縮檔）",
                        data=zip_buf.getvalue(),
                        file_name=f"專利說明書_{_get_client_now().strftime('%Y%m%d_%H%M')}.zip",
                        mime="application/zip",
                        type="primary",
                        use_container_width=True,
                    )
                with _btn_col2:
                    if st.button("🔄 重新查詢", key="btn_patent_reset", use_container_width=True):
                        _save_to_history()
                        _reset_patent_query()
                        st.rerun()
            else:
                # 沒有可下載檔案時，只顯示重新查詢
                if st.button("🔄 重新查詢", key="btn_patent_reset"):
                    _save_to_history()
                    _reset_patent_query()
                    st.rerun()

    # -- 查詢歷史紀錄 --
    _history = st.session_state.get('patent_history', [])
    if _history:
        st.divider()
        st.subheader("📜 查詢歷史")
        for _hi, _entry in enumerate(_history):
            _label = f"{_entry['timestamp']}（共 {_entry['total']} 筆：{_entry['ok_count']} 已下載 / {_entry['link_count']} 外國案連結）"
            with st.expander(_label, expanded=False):
                _h_results = _entry['results']
                _h_files = _entry['files']

                _h_ok = [r for r in _h_results if r['status'] == 'ok']
                _h_links = [r for r in _h_results if r['status'] == 'link']
                _h_not_found = [r for r in _h_results if r['status'] == 'not_found']
                _h_errors = [r for r in _h_results if r['status'] == 'error']

                if _h_ok:
                    st.markdown("**✅ 已下載：**")
                    for r in _h_ok:
                        st.markdown(f"- {r['raw']} → `{r['filename']}`")

                if _h_links:
                    st.markdown("**🔗 外國案連結：**")
                    for r in _h_links:
                        _pl = r.get("patent_links", [])
                        if _pl:
                            _lp = " ｜ ".join(f"[{lnk['source']}]({lnk['url']})" for lnk in _pl)
                            st.markdown(f"- {r['country']} {r['number']} → {_lp}")
                        else:
                            # 相容舊資料（只有 gpss_link）
                            st.markdown(f"- {r['country']} {r['number']} → [查詢]({r.get('gpss_link', '')})")

                if _h_not_found:
                    st.markdown("**⚠️ 未找到：**")
                    for r in _h_not_found:
                        st.markdown(f"- {r['raw']}：{r['error']}")

                if _h_errors:
                    st.markdown("**❌ 錯誤：**")
                    for r in _h_errors:
                        st.markdown(f"- {r['raw']}：{r['error']}")

                # 重新下載 ZIP
                if _h_files:
                    _zip_buf = BytesIO()
                    with zipfile.ZipFile(_zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for fname, fdata in _h_files.items():
                            zf.writestr(fname, fdata)
                    _zip_buf.seek(0)
                    st.download_button(
                        label=f"⬇️ 重新下載（{len(_h_files)} 個 PDF）",
                        data=_zip_buf.getvalue(),
                        file_name=f"專利說明書_{_entry['timestamp'].replace('-','').replace(':','').replace(' ','_')}.zip",
                        mime="application/zip",
                        key=f"btn_history_download_{_hi}",
                    )

# ============================================================
# 頁面 3: 設定
# ============================================================
elif _page == "⚙️ 設定":
    st.title("⚙️ 設定")

    st.subheader("台灣專利公開資訊 API")
    st.caption("由智慧財產局核發的 API 帳號密碼。每位使用者可各自儲存，帳密會加密保存。")

    _saved_creds = _load_api_credentials()
    _secrets_fallback = _load_secrets_fallback()

    _api_account = st.text_input(
        "API 帳號",
        value=_saved_creds['account'] if _saved_creds else "",
        key="tipo_account",
    )
    _api_password = st.text_input(
        "API 密碼",
        type="password",
        value=_saved_creds['password'] if _saved_creds else "",
        key="tipo_password",
    )

    _col_save, _col_spacer = st.columns([1, 3])
    with _col_save:
        if st.button("💾 儲存", type="primary", use_container_width=True):
            if _api_account.strip() and _api_password.strip():
                _save_api_credentials(_api_account.strip(), _api_password.strip())
                st.success("✅ 已儲存（加密）")
            else:
                st.warning("請輸入帳號和密碼")

    if _saved_creds:
        st.info("✅ 已有儲存的 API 帳號", icon="✅")
    elif _secrets_fallback:
        st.info("ℹ️ 尚未儲存個人 API 帳號，目前使用系統預設帳號。建議申請並填入自己的帳號。", icon="ℹ️")
    else:
        st.warning("⚠️ 尚未儲存 API 帳號，請輸入後點擊儲存。", icon="⚠️")

    st.divider()
    st.subheader("GPSS API")
    st.caption("全球專利檢索系統 API（需另外向智慧財產局申請 userCode）。取得後可於此填入，啟用外國專利自動下載功能。")
    st.info("🚧 功能開發中，待取得 userCode 後啟用。", icon="🚧")

# ============================================================
# 頁尾（顯示在 sidebar 底部）
# ============================================================
with st.sidebar:
    st.divider()
    if _GIT_COMMIT_UTC:
        if isinstance(_client_tz_offset, (int, float)):
            _client_tz = timezone(timedelta(minutes=-int(_client_tz_offset)))
            _last_update_local = _GIT_COMMIT_UTC.astimezone(_client_tz)
        else:
            _last_update_local = _GIT_COMMIT_UTC
        st.caption(f"{APP_VERSION} · {_last_update_local.strftime('%Y-%m-%d %H:%M')}")
    else:
        st.caption(f"{APP_VERSION}")
